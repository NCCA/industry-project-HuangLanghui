#!/usr/bin/env python
"""Convert one of the project's Markdown documents to a Word ``.docx`` file.

This is a lightweight, dependency-light Markdown-to-Word converter tailored to
the documents in ``report/`` (bilingual video script, coursework report). It is
*not* a full CommonMark implementation; it handles exactly the constructs those
files use: ATX headings, paragraphs, ``>`` block quotes, fenced code blocks,
pipe tables, ``-`` bullet lists, horizontal rules, and inline ``**bold**`` /
`` `code` `` spans. Chinese text renders natively in Word, so no font embedding
is required (unlike the HTML/PDF path).

Usage
-----
    python scripts/md_to_docx.py report/video_script_bilingual.md
    python scripts/md_to_docx.py report/video_script_bilingual.md report/out.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Author stamped into the document properties (overrides python-docx's default).
AUTHOR = "Langhui Huang"


def _add_inline(paragraph, text: str) -> None:
    """Add ``text`` to ``paragraph`` honouring **bold** and `code` inline spans."""
    # Split on **bold** and `code` while keeping the delimiters' content.
    tokens = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(tok)


def _shade(cell_or_para, hex_fill: str) -> None:
    """Apply a background shading to a table cell (used for the header row)."""
    tcPr = cell_or_para._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Fenced code block: gather until the closing ```.
        if line.lstrip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].lstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Pt(10)
            continue

        # Pipe table: consume consecutive table rows.
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].clear()
                _add_inline(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
                _shade(cell, "F0F0F0")
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[:len(header)]):
                    cells[j].paragraphs[0].clear()
                    _add_inline(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_inline(p, stripped[2:])
            for r in p.runs:
                r.italic = True
        elif re.match(r"^[-*] ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:])
        elif set(stripped) <= {"-"} and len(stripped) >= 3:
            doc.add_paragraph()  # horizontal rule -> spacing
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)

        i += 1

    # Set clean document properties (no tool boilerplate in author/comments).
    cp = doc.core_properties
    cp.author = AUTHOR
    cp.last_modified_by = AUTHOR
    cp.comments = ""

    doc.save(out_path)
    print(f"Wrote Word document: {out_path}")


def main() -> None:
    if len(sys.argv) < 2:
        print("usage: python scripts/md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    convert(md_path, out_path)


if __name__ == "__main__":
    main()
